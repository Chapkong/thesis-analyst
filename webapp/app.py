import json
import os
import re
from pathlib import Path

import pdfplumber
from anthropic import Anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from dotenv import load_dotenv
from flask import Flask, jsonify, request, send_file, send_from_directory

load_dotenv(Path(__file__).resolve().parent / ".env")

BASE_DIR = Path(__file__).resolve().parent.parent
INBOX_DIR = BASE_DIR / "inbox"
RESULTS_DIR = BASE_DIR / "results"
TEMPLATE_PATH = BASE_DIR / "templates" / "result_template.html"
AGENTS_DIR = BASE_DIR / ".claude" / "agents"

INBOX_DIR.mkdir(exist_ok=True)
RESULTS_DIR.mkdir(exist_ok=True)

MODEL = os.environ.get("CLAUDE_MODEL", "claude-opus-4-8")
FALLBACK_MODEL = os.environ.get("CLAUDE_FALLBACK_MODEL", "claude-sonnet-5")
client = Anthropic()

app = Flask(__name__, static_folder="static")

SECTION_KEYS = [
    "subjects",
    "questions_hypotheses",
    "data",
    "methodology",
    "findings_implications",
]

SECTION_TITLES = [
    ("subjects", "1. 연구대상"),
    ("questions_hypotheses", "2. 연구질문과 가설"),
    ("data", "3. 연구 데이터"),
    ("methodology", "4. 연구방법론"),
    ("findings_implications", "5. 주요결과 및 시사점"),
]


def load_agent_prompt(name):
    text = (AGENTS_DIR / f"{name}.md").read_text(encoding="utf-8")
    return text.split("---", 2)[-1].strip()


READER_PROMPT = load_agent_prompt("pdf-reader")
EXTRACTOR_PROMPT = load_agent_prompt("extractor")
VERIFIER_PROMPT = load_agent_prompt("verifier")


def extract_pdf_text(pdf_path):
    pages = []
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            pages.append(f"[p.{i}]\n{text}")
    return "\n\n".join(pages)


def _create_message(model, system, user, max_tokens, effort):
    return client.messages.create(
        model=model,
        max_tokens=max_tokens,
        system=system,
        thinking={"type": "adaptive"},
        output_config={"effort": effort},
        messages=[{"role": "user", "content": user}],
    )


def call_claude(system, user, max_tokens, effort="high"):
    resp = _create_message(MODEL, system, user, max_tokens, effort)
    if resp.stop_reason == "refusal" and FALLBACK_MODEL != MODEL:
        resp = _create_message(FALLBACK_MODEL, system, user, max_tokens, effort)
    text = "".join(block.text for block in resp.content if block.type == "text")
    if not text.strip():
        raise ValueError(
            f"모델 응답이 비어 있습니다 (stop_reason={resp.stop_reason}). "
            "문서가 너무 길어 max_tokens 안에서 답변을 끝맺지 못했거나 안전 필터에 의해 거부되었을 수 있습니다."
        )
    return text


def parse_json_block(text):
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if not match:
        raise ValueError("extractor 출력에서 JSON을 찾을 수 없습니다: " + text[:300])
    data = json.loads(match.group(0))
    for key in SECTION_KEYS:
        data.setdefault(key, ["○ 논문에 명시되지 않음"])
    header = data.get("header") or {}
    data["header"] = {
        "title": header.get("title") or "",
        "authors": header.get("authors") or "",
        "year": header.get("year") or "",
        "doc_type": header.get("doc_type") or "",
    }
    return data


def run_pipeline(pdf_path):
    raw_text = extract_pdf_text(pdf_path)
    if not raw_text.strip():
        raise ValueError("PDF에서 텍스트를 추출할 수 없습니다 (스캔본일 가능성이 높습니다)")

    reader_input = (
        "아래는 pdfplumber로 추출한 PDF 원문입니다 (페이지 마커 [p.N] 포함). "
        "이 텍스트를 바탕으로 지시된 구조로 정리하세요.\n\n" + raw_text
    )
    structured = call_claude(READER_PROMPT, reader_input, max_tokens=16000, effort="high")

    data = None
    fix_note = ""
    for attempt in range(3):
        extractor_input = structured
        if fix_note:
            extractor_input += "\n\n[verifier 수정 지시]\n" + fix_note
        extraction = call_claude(EXTRACTOR_PROMPT, extractor_input, max_tokens=16000, effort="high")
        data = parse_json_block(extraction)

        verify_input = (
            "[원문]\n" + structured +
            "\n\n[추출 결과 JSON]\n" + json.dumps(data, ensure_ascii=False, indent=2)
        )
        verdict = call_claude(VERIFIER_PROMPT, verify_input, max_tokens=4000, effort="medium")
        first_line = verdict.strip().splitlines()[0] if verdict.strip() else ""
        if "PASS" in first_line:
            break
        fix_note = verdict
    return data


def render_result_html(data):
    template = TEMPLATE_PATH.read_text(encoding="utf-8")
    payload = json.dumps(data, ensure_ascii=False).replace("</script>", "<\\/script>")
    return template.replace("/*__DATA__*/null", payload)


def render_result_docx(data, out_path):
    header = data.get("header") or {}
    doc = Document()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(header.get("title") or "논문 분석 결과")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    meta_parts = [p for p in [header.get("authors"), header.get("year"), header.get("doc_type")] if p]
    if meta_parts:
        meta_p = doc.add_paragraph(" · ".join(meta_parts))
        meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in meta_p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for key, title in SECTION_TITLES:
        heading_p = doc.add_paragraph()
        heading_p.paragraph_format.space_before = Pt(10)
        heading_run = heading_p.add_run(title)
        heading_run.bold = True
        heading_run.font.size = Pt(12)
        heading_run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

        for raw in data.get(key) or ["○ 논문에 명시되지 않음"]:
            t = str(raw).strip()
            is_sub = t.startswith("-")
            marker = "•" if is_sub else "○"
            content = t[1:].strip() if (is_sub or t.startswith("○")) else t
            item_p = doc.add_paragraph()
            item_p.paragraph_format.left_indent = Pt(28 if is_sub else 14)
            item_run = item_p.add_run(f"{marker} {content}")
            item_run.font.size = Pt(10.5)

    doc.save(out_path)


def safe_stem(filename):
    stem = Path(filename).stem
    stem = re.sub(r'[\\/:*?"<>|]', "_", stem)
    return stem or "result"


@app.route("/")
def index():
    return send_from_directory(app.static_folder, "index.html")


@app.route("/analyze", methods=["POST"])
def analyze():
    file = request.files.get("pdf")
    if not file or not file.filename.lower().endswith(".pdf"):
        return jsonify({"error": "PDF 파일을 선택하세요"}), 400

    stem = safe_stem(file.filename)
    pdf_path = INBOX_DIR / f"{stem}.pdf"
    file.save(pdf_path)

    try:
        data = run_pipeline(pdf_path)
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500

    out_html = f"{stem}.html"
    (RESULTS_DIR / out_html).write_text(render_result_html(data), encoding="utf-8")

    out_docx = f"{stem}.docx"
    render_result_docx(data, RESULTS_DIR / out_docx)

    return jsonify({"result": out_html, "docx": out_docx})


@app.route("/results/<path:filename>")
def view_result(filename):
    return send_from_directory(RESULTS_DIR, filename)


@app.route("/download/<path:filename>")
def download_result(filename):
    path = RESULTS_DIR / filename
    mimetype = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if path.suffix == ".docx"
        else None
    )
    return send_file(path, as_attachment=True, mimetype=mimetype)


if __name__ == "__main__":
    import threading
    import webbrowser

    threading.Timer(1.5, lambda: webbrowser.open("http://127.0.0.1:5000")).start()
    app.run(debug=False, port=5000)
